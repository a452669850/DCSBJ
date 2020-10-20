#!/usr/bin/env python
# -*- coding: utf-8 -*-

from PyQt5.QtCore import QThread, QWaitCondition, QMutex, pyqtSignal
from PyQt5.QtWidgets import QWidget, QVBoxLayout, QPushButton, QProgressBar, QMessageBox
from PyQt5.QtGui import QStandardItem
from utils.ClientModels import RunResult, Usecase, UsecaseGroup, Phrase, Procedure
from utils.WorkModels import PointModel, NetworkConfig
from utils.core import MainWindowConfig
from DcsUi.ExcelDockWidget import clearAllItem
from time import strftime
from pubsub import pub
import numpy
import uuid
import re
import json


class ProcedureThread(QThread):

    colorChange = pyqtSignal(int,int)
    logChange = pyqtSignal(int)
    quitThread = pyqtSignal()

    def __init__(self, *args, **kwargs):
        super(ProcedureThread, self).__init__(*args, **kwargs)
        self.pharses = Phrase.get_all()
        self.phrase = {}
        for phrase in self.pharses:
            phraseList = [phrase.id, phrase.name, phrase.operation]
            self.phrase[phrase.name] = phrase.operation       

        self.parent = kwargs['parent']

        # self.run_debug = True  # 初始的单步运行状态
        self.run_procedure_number = ""  # 当前运行的规程编号
        self.run_procedure_name = ""  # 当前运行的规程名称
        self.run_usecase_group_name = ""  # 当前运行的用例组名字
        self.run_usecase_number = ""  # 当前运行的用例编号
        self.run_operation_section = 0  # 当前规程运行的用例中具体段落的索引
        self.run_text = {}  # 运行中返回来的文本
        self.run_big_sort_index = 0  # 运行的单步操作的index，比如四个检查只一个单步操作
        self.run_sort = 0  # 运行的单行操作的sort，每个检查都有自己sort
        self.run_result = True  # 运行结果True 或者False
        self.run_uuid = uuid.uuid1()  # 运行时UID
        self.run_type_dict = {'procedure' : 1,
                        'usecase': 3,
                        'usecasegroup': 2}
        self.start_delay_dict = {}
        self.current_loop_list = []
        self.run_usecase_index = 0
        self.run_type_map = 0
        self.certification = {}
        self.typ = None
        
        self.is_stop = 0
        self._isPause = False
        self._isWork = False
        self.cond = QWaitCondition()
        self.mutex = QMutex()
        self.logChange.connect(self.parent.dockBottom.updateLog)
        self.tList = []
        self.atList = []
        self.shou = 0

    def pause(self):
        self._isPause = True

    def resume(self):
        self._isPause = False
        self.cond.wakeAll()

    def run(self):
        # 多线程执行
        self.procedureExcel = self.parent.dockTop.ExcelTab.currentWidget().proListView
        if self.parent.runCountEdit.text():
            self.runCount = int(self.parent.runCountEdit.text())
        else:
            self.runCount = 1
        if self.procedureExcel.model.item(0,5).text() == '高速':
            self.proType = '高速'
        elif self.procedureExcel.model.item(0,5).text() == '低速':
            self.proType = '低速'
        # print(self.runCount)
        if self.proType == '高速': 
            MainWindowConfig.IOMapping.stop_gather()
            self.msleep(1000)
            MainWindowConfig.IOMapping.START_GS_experiment()

        self._isWork = True
        clearAllItem(self.procedureExcel)
        self.colorChange.connect(self.procedureExcel.changeRowColor)
        if self.run_type == 'procedure':
            self.run_procedure_number = self.procedureExcel.model.item(0,3).text()
            self.run_procedure_name = self.procedureExcel.model.item(0,1).text()
            self.run_usecase_number = self.procedureExcel.model.item(1,3).text()
            self.judgeIndex = [5, 6]
        if self.run_type == 'usecase':
            self.run_usecase_number = self.procedureExcel.model.item(0,3).text()
            self.judgeIndex = [6, 6]
        if self.run_type == 'usecasegroup':
            self.run_usecase_group_name = self.procedureExcel.groupName
            self.judgeIndex = [6, 6]
        # print(self.runCount, 99999999999999999999999999)
        for i in range(self.runCount):
            clearAllItem(self.procedureExcel)
            while self._isWork:
                self.mutex.lock()
                if self._isPause:
                    self.cond.wait(self.mutex)
                # print(self.parent.procedureRunIndex, self.procedureExcel.colsLen)
                if self.parent.procedureRunIndex > self.procedureExcel.colsLen - 1:
                    self.run_result = True
                    if self.procedureExcel.model.item(0,5).text() == '高速':
                        #result = MainWindowConfig.IOMapping.receive_data()
                        #self.tList+=[float(x) for x in list(result)]
                        self.atList.append(self.tList)
                        self.tList = []
                        #print(result)
                    self.save_run_result()
                    self.msleep(MainWindowConfig.RunInterval)
                    self.mutex.unlock()
                    self.parent.procedureRunIndex = 0
                    
                    # reply = QMessageBox.question(self.parent, '提示', '执行完毕！', QMessageBox.Yes)
                    break
                    
                res = self.performAction(self.procedureExcel.getRowContent(self.parent.procedureRunIndex), self.procedureExcel.model.item(self.parent.procedureRunIndex,3).text())
                self.colorChange.emit(self.parent.procedureRunIndex, res)
                self.logChange.emit(self.parent.procedureRunIndex)
                self.parent.procedureRunIndex += 1
                self.msleep(500)
                self.mutex.unlock()
            self.parent.procedureRunIndex = 0
        # self.colorChange.emit(-self.parent.procedureRunIndex + 1)
        self._isWork = None
        self.parent.procedureRunIndex = 0
        if self.proType == '高速':
            #MainWindowConfig.IOMapping.STOP_GS_experiment()
            #self.msleep(1000)
            #MainWindowConfig.IOMapping.start_gather(self.parent)
            #wordPath = Procedure.get(name = self.parent.currentPro).wordPath
            #insertDocx(wordPath, self.atList)
            self.endHighPro()
        self.atList = []
        self.exec_()
    
    def endHighPro(self):
        MainWindowConfig.IOMapping.STOP_GS_experiment()
        self.msleep(1000)
        MainWindowConfig.IOMapping.start_gather(self.parent)
        wordPath = Procedure.get(name = self.parent.currentPro).wordPath
        insertDocx(wordPath, self.atList)

    def save_run_result(self):
        # 保存记录
        self.run_type_map = self.run_type_dict[self.run_type]
        runresult = RunResult()
        runresult.run_uuid = self.run_uuid
        runresult.procedure_number = self.run_procedure_number
        runresult.procedure_name = self.run_procedure_name
        runresult.usecase_group_name = self.run_usecase_group_name
        runresult.usecase_number = self.run_usecase_number
        runresult.operation_section = self.run_operation_section
        runresult.run_big_sort = self.run_big_sort_index
        runresult.section_sort = self.run_sort
        runresult.run_text = self.run_text
        runresult.run_result = self.run_result
        runresult.run_time = strftime("%Y-%m-%d %H:%M:%S")
        runresult.run_type = self.run_type_map
        runresult.run_usecase_index = self.run_usecase_index
        # 如果usecase_number 、operation_section、section_sort都相同，则是循环的操作，
        # 只需要更新即可，不能用保存
        exist_runresult = RunResult.select().where(RunResult.usecase_number == runresult.usecase_number,
                                                   RunResult.usecase_group_name == runresult.usecase_group_name,
                                                   # RunResult.run_uuid == runresult.run_uuid,
                                                   RunResult.operation_section == runresult.operation_section,
                                                   RunResult.section_sort == runresult.section_sort)
        if exist_runresult.first() == None:
            runresult.is_stop = self.is_stop
            runresult.run_text = json.dumps(self.run_text) 
            runresult.run_result = self.run_result
            runresult.run_time = strftime("%Y-%m-%d %H:%M:%S")
            runresult.certification = json.dumps(self.certification)
            runresult.save()
        if exist_runresult.first() != None:
            update_result = exist_runresult.first()
            update_result.run_text = json.dumps(self.run_text)
            update_result.run_result = self.run_result
            update_result.run_time = strftime("%Y-%m-%d %H:%M:%S")
            update_result.is_stop = self.is_stop
            update_result.certification = json.dumps(self.certification)
            RunResult.update_obj(update_result)

    def performAction(self, action, typ):
        # 步骤执行函数
        if action is None or action == ' ':
            return
        if self.run_type == 'procedure':
            exResIndex = 3
            REsIndex = 4
            timeIndex = 7
        else:
            exResIndex = 4
            REsIndex = 4
            timeIndex = 2
        if '等待' in action:
            # print(int(re.findall('\d+',action)[0])*1000)
            if self.typ:
                MainWindowConfig.IOMapping.High_speed_test(int(float(self.typ)), 7)
            else:
                MainWindowConfig.IOMapping.doWrite()
            self.msleep(int(re.findall('\d+',action)[0])*1000)
            self.run_text[action] = ['是', strftime("%Y-%m-%d %H:%M:%S"), '成功']
            return 1
                
        if '接收' in action:
            MainWindowConfig.IOMapping.High_speed_test(int(float(self.typ)), 7)
            result = MainWindowConfig.IOMapping.receive_data()
            self.tList+=[float(x)/1000 for x in list(result)]
            # print(result, 7777777777777777777777777777)
            self.typ = None
            self.run_text[action] = ['是', strftime("%Y-%m-%d %H:%M:%S"), '成功']
            return 1

        if '收到' in action:
            # print(self.shou)
            MainWindowConfig.IOMapping.doWrite()
            #     self.shou = 1
            #     self.resume()
            #     self.msleep(300)
            read = action.split('\n')[1].split('=')[0]
            slep = 1
            highValue = PointModel.get(PointModel.sig_name == read).rhi
            lowValue = PointModel.get(PointModel.sig_name == read).rlo
            while slep:
                if self._isPause:
                    MainWindowConfig.IOMapping.readall()
                if slep >= 15001:
                    self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[0], QStandardItem('否'))
                    self.run_text[action] = ['否', strftime("%Y-%m-%d %H:%M:%S"), '失败']
                    self.parent.log.infoLog(f'执行到{self.parent.procedureRunIndex}行,{action}失败')
                    # if self.shou == 1:
                    #     print(33333333333)
                    #     self.msleep(300)
                    #     # self.is_stop = 0
                    #     # self.cond.wait(self.mutex)
                    #     self._isPause = True
                    #     self.pause()
                    #     # self.stop()
                        # 
                    return 0
                    break
                result = float(MainWindowConfig.IOMapping.read(read))
                expectedResult = float(action.split('\n')[1].split('=')[1])
                self.procedureExcel.model.setItem(self.parent.procedureRunIndex, REsIndex, QStandardItem(read + '=' + str(result)))
                # Operator = self.getOperator(re.compile('[<,<=,>,>=,=,≤,≥,==]').search(action))
                # print(result, expectedResult)
                if highValue or lowValue:
                    highValue = float(highValue)
                    lowValue = float(lowValue)
                    # print(expectedResult - 0.02 * (highValue - lowValue))
                    # print(expectedResult + 0.02 * (highValue - lowValue))
                    if result >= expectedResult - 0.02 * (highValue - lowValue)  and result <= expectedResult + 0.02 * (highValue - lowValue):
                        self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[0], QStandardItem('是'))
                        self.run_text[action] = ['是', strftime("%Y-%m-%d %H:%M:%S"), '成功']
                        self.parent.log.infoLog(f'执行到{self.parent.procedureRunIndex}行,{action}成功')
                        return 1
                        # if self.shou == 1:
                        #     self.msleep(300)
                        #     # self.is_stop = 0
                        #     # self.cond.wait(self.mutex)
                        #     self._isPause = True
                        #     self.pause()
                        # self.certification[action] = f'解析{action},通过短语库字段检索,判定该步骤为{value},解析到需要操作的变量为{read},{self.search(read, value, result)},与预期结果{read}={expectedResult}相同'
                        #self.certification[action] = [value,read,self.search(read),result,True]
                        # #self.certification[action] = f'解析{action},通过短语库字段检索,判定该步骤为{value},解析到需要操作的变量为{read},{self.search(read)},与预期结果{read}={expectedResult}不符'
                        #self.certification[action] = [value,read,self.search(read),result,False]
                        # return 0
                else:
                    if expectedResult == result:
                        self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[0], QStandardItem('是'))
                        # self.run_text[action] = ['是', strftime("%Y-%m-%d %H:%M:%S"), read + '=' + str(result)]
                        self.parent.log.infoLog(f'执行到{self.parent.procedureRunIndex}行,{action}成功')
                        self.run_text[action] = ['是', strftime("%Y-%m-%d %H:%M:%S"), '成功']
                        # if self.shou == 1:
                        #     self.msleep(300)
                        #     # self.is_stop = 0
                        #     # self.cond.wait(self.mutex)
                        #     self._isPause = True
                        #     self.pause()
                        return 1
                slep += 500
                self.msleep(500)
            # else:
            # print(result, 2222222)

            return
            
        for key, value in self.phrase.items():
            if key in action:
                if value == 'SET' or not re.compile(u'[\u4e00-\u9fa5]').search(action):
                    # print(key)
                    if key in action: 
                        writeList = action.split('\n')[1:]
                    else:
                        writeList = action.split('\n')
                    for write in writeList:
                        write = write.split('=')
                        print(write)
                        try:
                            if re.findall('\d+',typ):
                                typ = int(float(typ))
                                print(write[0], 11111)
                                self.typ = typ
                                self.procedureExcel.model.setItem(self.parent.procedureRunIndex, timeIndex, QStandardItem(strftime("%Y-%m-%d %H:%M:%S")))
                                result = MainWindowConfig.IOMapping.GS_write(write[0], float(re.findall('\d+',write[1])[0]), int(float(typ)))
                                pub.subscribe(snoop, pub.ALL_TOPICS)
                                self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[0], QStandardItem('是'))
                                self.run_text[action] = ['是', strftime("%Y-%m-%d %H:%M:%S"), '成功']
                                self.parent.log.infoLog(f'执行到{self.parent.procedureRunIndex}行,{action}成功')
                                #self.certification[action] = [value,write[0],self.search(write[0]), result]
                                return 1
                            else:
                                self.procedureExcel.model.setItem(self.parent.procedureRunIndex, timeIndex, QStandardItem(strftime("%Y-%m-%d %H:%M:%S")))
                                result = MainWindowConfig.IOMapping.runWrite(write[0], float(write[1]))
                                pub.subscribe(snoop, pub.ALL_TOPICS)
                                # print(MainWindowConfig.IOMapping.write(write[0], float(re.findall('\d+',write[1])[0])))
                                # if result:
                                self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[0], QStandardItem('是'))
                                self.procedureExcel.model.setItem(self.parent.procedureRunIndex, REsIndex, QStandardItem(write[0] + '=' + str(result)))
                                self.run_text[action] = ['是', strftime("%Y-%m-%d %H:%M:%S"), '成功']
                                self.parent.log.infoLog(f'执行到{self.parent.procedureRunIndex}行,{action}成功')
                                # #self.certification[action] = f'解析{action},通过短语库字段检索,判定该步骤为{value},解析到需要操作的变量为{write[0]},{self.search(write[0], value, result[0])}'
                                #self.certification[action] = [value,write[0],self.search(write[0]), result]
                                return 1
                                # else:
                                    
                                #     self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[1], QStandardItem('否'))
                                #     self.run_text[action] = ['否', strftime("%Y-%m-%d %H:%M:%S")]
                                #     self.parent.log.infoLog(f'执行到{self.parent.procedureRunIndex}行,{action}失败')
                                #     return 0
                        except Exception as e:
                            error = self.judgeError(e, write[0])
                            self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[1], QStandardItem('否'))
                            self.run_text[action] = ['否', strftime("%Y-%m-%d %H:%M:%S"),error]
                            self.procedureExcel.model.setItem(self.parent.procedureRunIndex, REsIndex, QStandardItem(str(error)))
                            self.parent.log.warningLog(f'执行到{self.parent.procedureRunIndex}行,{action}失败,错误原因{e}')
                            # #self.certification[action] = f'解析{action},通过短语库字段检索,判定该步骤为{value},解析到需要操作的变量为{write[0]},设置失败,错误原因{str(error)}'
                            #self.certification[action] = [value,write[0],str(error)]
                            return 0
                elif value in 'CHECK':
                    read = re.sub('[\u4e00-\u9fa5]', '', action)
                    try:
                        self.procedureExcel.model.setItem(self.parent.procedureRunIndex, timeIndex, QStandardItem(strftime("%Y-%m-%d %H:%M:%S")))
                        result = str(MainWindowConfig.IOMapping.read(read))
                        pub.subscribe(snoop, pub.ALL_TOPICS)
                    except Exception as e:
                        error = self.judgeError(e, read)
                        self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[0], QStandardItem('否'))
                        self.run_text[action] = ['否', strftime("%Y-%m-%d %H:%M:%S"), error]
                        self.procedureExcel.model.setItem(self.parent.procedureRunIndex, REsIndex, QStandardItem(str(error)))
                        self.parent.log.warningLog(f'执行到{self.parent.procedureRunIndex}行,{action}失败,错误原因{e}')
                        # #self.certification[action] = f'解析{action},通过短语库字段检索,判定该步骤为{value},解析到需要操作的变量为{read},读取失败,错误原因为{str(error)}'
                        #self.certification[action] = [value,read,str(error)]
                        result = None
                        return 0
                    expectedResult = self.procedureExcel.model.item(self.parent.procedureRunIndex, exResIndex).text().split('=')[-1]
                    self.procedureExcel.model.setItem(self.parent.procedureRunIndex, REsIndex, QStandardItem(read + '=' + result))
                    Operator = self.getOperator(re.compile('[<,<=,>,>=,=,≤,≥,==]').search(action))
                    if eval(f'{str(float(expectedResult))}{Operator}{str(float(result))}'):
                        self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[0], QStandardItem('是'))
                        self.run_text[action] = ['是', strftime("%Y-%m-%d %H:%M:%S"), read + '=' + result]
                        self.parent.log.infoLog(f'执行到{self.parent.procedureRunIndex}行,{action}成功')
                        # #self.certification[action] = f'解析{action},通过短语库字段检索,判定该步骤为{value},解析到需要操作的变量为{read},{self.search(read, value, result)},与预期结果{read}={expectedResult}相同'
                        #self.certification[action] = [value,read,self.search(read),result,True]
                        return 1
                    else:
                        self.procedureExcel.model.setItem(self.parent.procedureRunIndex, self.judgeIndex[1], QStandardItem('否'))
                        self.run_text[action] = ['否', strftime("%Y-%m-%d %H:%M:%S"), read + '=' + result]
                        self.parent.log.infoLog(f'执行到{self.parent.procedureRunIndex}行,{action}成功')
                        # #self.certification[action] = f'解析{action},通过短语库字段检索,判定该步骤为{value},解析到需要操作的变量为{read},{self.search(read)},与预期结果{read}={expectedResult}不符'
                        #self.certification[action] = [value,read,self.search(read),result,False]
                        return 0

    def judgeError(self, e, value):
        if str(e) == value:
            error = '未在变量表找到该变量'
        elif str(e) == '1':
            error = '功能码异常'
        elif str(e) == '2':
            error = '地址异常'
        elif str(e) == '[WinError 10061] 由于目标计算机积极拒绝，无法连接。':
            error = '未连接'
        else:
            error = str(e)
        return error

    def search(self, name):
        # 通道号
        channel = PointModel.get(PointModel.sig_name == name).channel
        # 通信接口
        slot = PointModel.get(PointModel.sig_name == name).slot
        # 协议
        protocol = NetworkConfig.get(NetworkConfig.slot == slot).protocol
        return [channel,protocol]

    def getOperator(self, Operator):
        if Operator == '=':
            return '=='
        elif Operator == '≤':
            return '<='
        elif Operator == '≥':
            return '>='
        else:
            return Operator

def snoop(topicObj=pub.AUTO_TOPIC, **mesgData):
    print('topic "%s": %s' % (topicObj.getName(), mesgData))

def average(lis):
    if len(lis) == 1:
        lis = [float(x) for x in lis[0]]
        return lis
    else:
        c = numpy.array(lis)
        # c = [float(x) for x in c]
        return c.mean(axis=0)

def insertDocx(wordPath, tList):
    # print(wordPath,tList)
    txtPath = wordPath.split('.')[0] + ' ' + strftime("%Y-%m-%d %H-%M-%S") + '.txt'
    f = open(txtPath, 'w')
    x = 1
    z = 1
    for t in tList:
        for y in t:
            f.write(' △T{} = '.format(x) + str(y))
            x += 1
            if x == 17:
                f.write('\r\n')
                x = 1
        f.write(str(z))
        f.write('\r\n')
        z += 1
    f.close()

    index = 0
    from docx import Document
    tList = average(tList)
    # 获取word文档
    doc = Document(wordPath)
    print(len(tList))
    # for row, t in zip(doc.tables[-1].rows[1:],tList):  # 遍历表格中的所有行
    #     res = row.cells[0:16][-1].text.split('=')[0] + '=' + str(t)
    #     row.cells[0:16][-1].text = res
        # print(row.cells[0:16][-1].text)
    for row in doc.tables[-1].rows[1:]:  # 遍历表格中的所有行
        for cell in row.cells[1:17]:
            cell.text = str(tList[index])
            index += 1
    index = 0
    doc.save(wordPath)

