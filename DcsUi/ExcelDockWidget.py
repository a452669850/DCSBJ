from PyQt5.QtWidgets import QTabWidget, QDockWidget, QTableView, QHeaderView, QTableWidget, QMessageBox,\
QAbstractItemView, QMenu, QWidget, QTextBrowser, QFrame, QGridLayout, QLabel, QWidget, QVBoxLayout, QMainWindow
from PyQt5.QtWebEngineWidgets import *
from PyQt5.QtGui import QStandardItemModel, QStandardItem, QBrush, QCursor, QColor
from PyQt5.QtCore import Qt, QCoreApplication, QMetaObject, QUrl, QFileInfo
from DcsUi.DockCLass import NewDockWidget
from utils.ClientModels import Usecase, UsecaseGroup, Procedure
from PyQt5.QAxContainer import QAxWidget
import openpyxl
import json
import sys
import os
import docx
from pydocx import PyDocX

tittle = ['序号', '实验步骤', '执行时间', '操作类型', '执行结果', '备注', '是否符合预期']

class TabDockWidget(NewDockWidget):
    def __init__(self, title, parent = None):
        NewDockWidget.__init__(self, title, parent=parent)
        self.ExcelTab = QTabWidget()
        self.ExcelTab.setTabsClosable(True)
        self.ExcelTab.tabCloseRequested.connect(self.closeTab)
        self.parent = parent

    # def addExcel(self, path):
    #     if path:

    #         self.TabelView = ExcelTabelView(path, self.parent)
    #         self.ExcelTab.addTab (self.TabelView, '规程：' + os.path.basename(path))
    #         self.setWidget(self.ExcelTab)
    
    def addExcel(self, path):
        if path:
            self.TabelView = ProWidget(path, self.parent)
            self.ExcelTab.addTab (self.TabelView, '规程：' + os.path.basename(path))
            self.setWidget(self.ExcelTab)

    def addUseCase(self, usecaseName):
        self.TabelView = UsecaseTabelView(usecaseName, self.parent)
        self.ExcelTab.addTab (self.TabelView, '用例：' + usecaseName)
        self.setWidget(self.ExcelTab)

    def addUsecaseGroup(self, allCases):
        self.TabelView = ProWidget(allCases, self.parent)
        self.ExcelTab.addTab (self.TabelView, 'TEST1')
        self.setWidget(self.ExcelTab)

    def closeTab(self, index):
        if not self.parent.ProcedureThread._isWork or index != self.ExcelTab.indexOf(self.parent.ProcedureThread.procedureExcel):
            self.ExcelTab.removeTab(index)
        else:
            reply = QMessageBox.question(self, '提示', '请先退出正在运行的规程！', QMessageBox.Yes)


class ExcelTabelView(QTableView):
    def __init__(self, path, parent):
        QTableView.__init__(self)
        self.parent = parent
        if os.path.exists(path):
            self.model = QStandardItemModel(0,0)
            self.addExcelContent(path)
            self.setModel(self.model)
            self.resizeRowsToContents()
            self.horizontalHeader().setStretchLastSection(True)
            self.setEditTriggers(QAbstractItemView.NoEditTriggers)
            self.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
            # self.horizontalHeader().setPropertySectionResizeMode(QHeaderView.ResizeToContents);
            self.verticalHeader().setSectionResizeMode(QHeaderView.Fixed);
            self.setContextMenuPolicy(Qt.CustomContextMenu) # 右键菜单
            self.customContextMenuRequested.connect(self.showContextMenu) # 菜单连接信号
            self.type = 'procedure'

            
    def addExcelContent(self,path):
        self.setUpdatesEnabled(False) 
        wb = openpyxl.load_workbook(path)
        ws = wb.active
        self.rowsCon = ws.rows
        for row in ws.rows:
            self.model.appendRow([QStandardItem(str(x.value)) if x.value is not None else QStandardItem(' ') for x in row])
        self.rowsLen = ws.max_column
        self.colsLen = ws.max_row
        self.rowsCon = [x for x in self.rowsCon]
        self.setUpdatesEnabled(True) 

    def changeRowColor(self, rowIndex, res):
        if res:
            brush = QBrush(QColor(0,255,154))
        else:
            brush = QBrush(QColor(220,20,60))
        if rowIndex < 0:
            for x in range(self.rowsLen):
                self.model.setData(self.model.index(abs(rowIndex), x), brush, Qt.BackgroundRole)
        else:
            for x in range(self.rowsLen):
                self.model.setData(self.model.index(rowIndex, x), brush, Qt.BackgroundRole)
                # self.model.setData(self.model.index(rowIndex -1, x), QBrush(Qt.gray), Qt.BackgroundRole)

    def getRowContent(self, rowIndex):
        try:
            rowCon =  self.model.item(rowIndex, 1).text()
        except:
            try:
                if 'STEP' in self.model.item(rowIndex, 1).text():
                    return None
            except:
                return None
            return '全部'
        return rowCon

    def showContextMenu(self):  # 创建右键菜单、
        self.contextMenu = QMenu(self)
        self.actionA = self.contextMenu.addAction('从选中行开始执行')
        self.contextMenu.popup(QCursor.pos())  # 2菜单显示的位置
        self.actionA.triggered.connect(self.actionHandler)
        self.contextMenu.show()

    def actionHandler(self):
        if self.parent.ProcedureThread._isPause and self.parent.ProcedureThread._isWork:
            self.parent.procedureRunIndex = self.currentIndex().row()
            self.parent.ProcedureThread.resume()
        else:
            self.parent.procedureRunIndex = self.currentIndex().row()
            self.parent.procedureAutoRunClicked()
        self.parent.log.infoLog(f'{self.parent.procedureRunPath}规程从{self.currentIndex().row()}行开始执行')

class UsecaseTabelView(QTableView):
    def __init__(self, usecaseName, parent):
        QTableView.__init__(self)
        self.parent = parent
        self.type = 'usecase'
        self.colIndex = 0
        self.rowIndex = 2
        self.usecaseName = usecaseName
        self.usecaseOperation = json.loads(Usecase.get_by_name(usecaseName).operation)
        self.model = QStandardItemModel(0,0)
        self.addExcelContent()
        self.setModel(self.model)
        self.resizeRowsToContents()
        self.horizontalHeader().setStretchLastSection(True)
        self.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        # self.horizontalHeader().setPropertySectionResizeMode(QHeaderView.ResizeToContents);
        self.verticalHeader().setSectionResizeMode(QHeaderView.Fixed);
        self.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.setContextMenuPolicy(Qt.CustomContextMenu) # 右键菜单
        self.customContextMenuRequested.connect(self.showContextMenu) # 菜单连接信号

            
    def addExcelContent(self): 
        self.model.setItem(0, 0, QStandardItem('测试用例'))
        self.model.setItem(0, 1, QStandardItem(self.usecaseName))
        self.model.setItem(0, 2, QStandardItem('用例编号'))
        self.model.setItem(0, 3, QStandardItem(Usecase.get_by_name(self.usecaseName).number))
        self.model.appendRow([QStandardItem(str(x)) if x is not None else QStandardItem(' ') for x in tittle])
        self.setUpdatesEnabled(False)
        for step in self.usecaseOperation:
            for opr in step:
                # print(opr)
                if 'STEP' in opr:
                    self.model.setItem(self.rowIndex, self.colIndex, QStandardItem(opr))
                    self.rowIndex += 1
                else:
                    # self.insertRow(self.rowIndex)
                    for k, v in opr[1].items():
                        v = '-' if v == '' else v
                        self.model.setItem(self.rowIndex, self.colIndex, QStandardItem(v))
                        self.colIndex += 1
                    self.rowIndex += 1
                    self.colIndex = 0
        self.colsLen = self.model.rowCount()
        self.rowsLen = 7
        self.setUpdatesEnabled(True) 

    def changeRowColor(self, rowIndex, res):
        if res:
            brush = QBrush(QColor(0,255,154))
        else:
            brush = QBrush(QColor(220,20,60))
        try:
            self.model.item(rowIndex, 1).text()
        except:
            # for x in range(self.rowsLen):
            #     self.model.setData(self.model.index(rowIndex -1, x), QBrush(Qt.gray), Qt.BackgroundRole)
            return
        if rowIndex < 0:
            for x in range(self.rowsLen):
                self.model.setData(self.model.index(abs(rowIndex), x), brush, Qt.BackgroundRole)
        else:
            # print([x.text() for x in self.model.takeRow(rowIndex)])
            for x in range(self.rowsLen):
                self.model.setData(self.model.index(rowIndex, x), brush, Qt.BackgroundRole)
                # self.model.setData(self.model.index(rowIndex -1, x), QBrush(Qt.gray), Qt.BackgroundRole)

    def getRowContent(self, rowIndex):
        try:
            rowCon =  self.model.item(rowIndex, 1).text()
        except:
            try:
                if 'STEP' in self.model.item(rowIndex, 1).text():
                    return None
            except:
                return None
            return '全部'
        return rowCon

    def showContextMenu(self):  # 创建右键菜单、
        self.contextMenu = QMenu(self)
        self.actionA = self.contextMenu.addAction('从选中行开始执行')
        self.contextMenu.popup(QCursor.pos())  # 2菜单显示的位置
        self.actionA.triggered.connect(self.actionHandler)
        self.contextMenu.show()

    def actionHandler(self):
        if self.parent.ProcedureThread._isPause and self.parent.ProcedureThread._isWork:
            self.parent.procedureRunIndex = self.currentIndex().row()
            self.parent.ProcedureThread.resume()
        else:
            self.parent.procedureRunIndex = self.currentIndex().row()
            self.parent.procedureAutoRunClicked()
        self.parent.log.infoLog(f'{self.parent.procedureRunPath}规程从{self.currentIndex().row()}行开始执行')

class UsecaseGroupView(QTableView):
    def __init__(self, allCases, parent):
        QTableView.__init__(self)
        self.parent = parent
        self.type = 'usecasegroup'
        self.allCases = allCases
        self.colIndex = 0
        self.rowIndex = 0
        self.model = QStandardItemModel(0,0)
        self.getAllUsecase()
        self.setModel(self.model)
        self.resizeRowsToContents()
        self.horizontalHeader().setStretchLastSection(True)
        self.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.verticalHeader().setSectionResizeMode(QHeaderView.Fixed)
        # self.horizontalHeader().setPropertySectionResizeMode(QHeaderView.ResizeToContents)
        # self.verticalHeader().setMinimumHeight(100)
        # self.verticalHeader().setMaximumHeight(100)
        self.setEditTriggers(QAbstractItemView.NoEditTriggers)

        self.setContextMenuPolicy(Qt.CustomContextMenu) # 右键菜单
        self.customContextMenuRequested.connect(self.showContextMenu) # 菜单连接信号
        self.proType = parent.proType

    def getAllUsecase(self):
        for usecaseName in self.allCases:
            self.usecaseOperation = json.loads(Usecase.get_by_name(usecaseName).operation)
            self.insertUsecase(usecaseName)

    def insertUsecase(self, usecaseName):
        self.setUpdatesEnabled(False)
        self.model.setItem(self.rowIndex, 0, QStandardItem('测试用例'))
        self.model.setItem(self.rowIndex, 1, QStandardItem(usecaseName))
        self.model.setItem(self.rowIndex, 2, QStandardItem('用例编号'))
        self.model.setItem(self.rowIndex, 3, QStandardItem(Usecase.get_by_name(usecaseName).number))
        self.model.setItem(self.rowIndex, 4, QStandardItem('实验类型'))
        self.model.setItem(self.rowIndex, 5, QStandardItem(self.parent.proType))
        self.model.appendRow([QStandardItem(str(x)) if x is not None else QStandardItem(' ') for x in tittle])
        self.rowIndex += 2
        for step in self.usecaseOperation:
            for opr in step:
                # print(opr)
                if 'STEP' in opr[0]:
                    self.model.setItem(self.rowIndex, self.colIndex, QStandardItem(opr[0]))
                    self.model.setItem(self.rowIndex, self.colIndex + 1, QStandardItem(opr[1]))
                    self.rowIndex += 1
                elif '序号' in opr[0]:
                    self.model.setItem(self.rowIndex, self.colIndex, QStandardItem(''))
                    self.model.setItem(self.rowIndex + 1, self.colIndex, QStandardItem(usecaseName))
                    self.rowIndex += 2
                elif opr[0]:
                    # print(opr)
                    for k, v in opr[1].items():
                        # print(v)
                        self.model.setItem(self.rowIndex, self.colIndex, QStandardItem(str(v)))
                        self.colIndex += 1
                    self.rowIndex += 1
                    self.colIndex = 0
        self.rowIndex += 1
        self.colsLen = self.model.rowCount()
        self.rowsLen = 7
        self.setUpdatesEnabled(True) 

    def changeRowColor(self, rowIndex, res):
        if res:
            brush = QBrush(QColor(0,255,154))
        else:
            brush = QBrush(QColor(220,20,60))
        try:
            self.model.item(rowIndex, 1).text()
        except:
            # for x in range(self.rowsLen):
            #     self.model.setData(self.model.index(rowIndex -1, x), QBrush(Qt.gray), Qt.BackgroundRole)
            return
        if rowIndex < 0:
            for x in range(self.rowsLen):
                self.model.setData(self.model.index(abs(rowIndex), x), brush, Qt.BackgroundRole)
        else:
            # print([x.text() for x in self.model.takeRow(rowIndex)])
            for x in range(self.rowsLen):
                self.model.setData(self.model.index(rowIndex, x), brush, Qt.BackgroundRole)
                # self.model.setData(self.model.index(rowIndex -1, x), QBrush(Qt.gray), Qt.BackgroundRole)

    def getRowContent(self, rowIndex):
        try:
            rowCon =  self.model.item(rowIndex, 1).text()
            # print(rowCon,1111111111111111)
        except:
            try:
                if 'STEP' in self.model.item(rowIndex, 1).text():
                    return None
            except:
                return None
            return '全部'
        return rowCon

    def showContextMenu(self):  # 创建右键菜单、
        self.contextMenu = QMenu(self)
        self.actionA = self.contextMenu.addAction('从选中行开始执行')
        self.contextMenu.popup(QCursor.pos())  # 2菜单显示的位置
        self.actionA.triggered.connect(self.actionHandler)
        self.contextMenu.show()

    def actionHandler(self):
        if self.parent.ProcedureThread._isPause and self.parent.ProcedureThread._isWork:
            self.parent.procedureRunIndex = self.currentIndex().row()
            self.parent.ProcedureThread.resume()
        else:
            self.parent.procedureRunIndex = self.currentIndex().row()
            self.parent.procedureAutoRunClicked()
        self.parent.log.infoLog(f'{self.parent.procedureRunPath}规程从{self.currentIndex().row()}行开始执行')

def clearAllItem(tabelview):
    for y in range(tabelview.model.rowCount()):
        for x in range(tabelview.model.columnCount()):
            tabelview.model.setData(tabelview.model.index(y, x), QBrush(), Qt.BackgroundRole)

class ProWidget(QWidget):
    def __init__(self, path, parent): 
        super().__init__()
        self.path = path
        self.parent = parent
        self.setupUi()

    def setupUi(self):
        self.gridLayout = QGridLayout(self)
        self.gridLayout.setObjectName("gridLayout")
        self.label_2 = QLabel(self)
        self.label_2.setAlignment(Qt.AlignCenter)
        self.label_2.setObjectName("label_2")
        self.gridLayout.addWidget(self.label_2, 0, 2, 1, 1)
        self.wordPath = Procedure.get(name = self.parent.currentPro).wordPath
        self.wordAx = AxWidget(self.wordPath, self.parent)
        self.wordAx.setObjectName("textBrowser")
        # self.wordAx.show()
        self.gridLayout.addWidget(self.wordAx, 1, 2, 1, 1)
        self.proListView = UsecaseGroupView(self.path, self.parent)
        self.proListView.setObjectName("proListView")
        self.gridLayout.addWidget(self.proListView, 1, 0, 1, 1)
        self.label = QLabel(self)
        self.label.setAlignment(Qt.AlignCenter)
        self.label.setObjectName("label")
        self.gridLayout.addWidget(self.label, 0, 0, 1, 1)
        self.line = QFrame(self)
        self.line.setFrameShape(QFrame.VLine)
        self.line.setFrameShadow(QFrame.Sunken)
        self.line.setObjectName("line")
        self.gridLayout.addWidget(self.line, 0, 1, 2, 1)
        self.gridLayout.setColumnStretch(0, 3)
        self.gridLayout.setColumnStretch(1, 1)
        self.gridLayout.setColumnStretch(2, 1)
        self.retranslateUi()
        QMetaObject.connectSlotsByName(self)

    # def addWord(self):


    def retranslateUi(self):
        _translate = QCoreApplication.translate
        self.setWindowTitle(_translate("self", "self"))
        self.label_2.setText(_translate("self", "规程源文件"))
        self.label.setText(_translate("self", "解析后规程"))

class AxWidget(QMainWindow):
 
    def __init__(self,path,parent, *args, **kwargs):
        super(AxWidget, self).__init__(*args, **kwargs)
        # self.resize(800, 600)
        # layout = QVBoxLayout(self)
        # self.axWidget = QAxWidget(self)
        # self.axWidget = QTextBrowser()
        # self.setCentralWidget(self.axWidget)
        # path = 'C:\\Users\\lj\\Desktop\\原量程中子注量率高 分析(1).docx'
        # my_doc = docx.Document(path)
            # print(my_doc.paragraphs[0].text)
        # for my_paragraph in my_doc.paragraphs:
        #     print(my_paragraph.text)
        #     self.axWidget.append(my_paragraph.text)
        # self.axWidget.append(read_docx(path))

        # layout.addWidget(self.axWidget)
        # layout.addWidget(QLabel())
        # self.openOffice(path)
        # from pydocx import PyDocX
        html = PyDocX.to_html(path)
        docxName = os.path.basename(path)
        portion = os.path.splitext(docxName)
        htmlName = portion[0] + '.html'
        htmlPath = os.path.join(parent.projectPath, '规程文档', htmlName)
        f = open(htmlPath, 'w', encoding="utf-8")
        f.write(html)
        f.close()

        # self.parent = parent
        
        self.browser=QWebEngineView()
        #加载外部的web界面
        self.browser.load(QUrl(QFileInfo(htmlPath).absoluteFilePath()))
        self.setCentralWidget(self.browser)

        
        
        
 
    def openOffice(self, path):
        app = 'Word.Application'
        self.axWidget.clear()
        if not self.axWidget.setControl(app):
            return QMessageBox.critical(self, '错误', '没有安装  %s' % app)
        self.axWidget.dynamicCall(
            'SetVisible (bool Visible)', 'false')  # 不显示窗体
        self.axWidget.setProperty('DisplayAlerts', False)
        self.axWidget.setControl(path)
        self.axWidget.show()
 
    def openPdf(self, path):
        self.axWidget.clear()
        if not self.axWidget.setControl('Adobe PDF Reader'):
            return QMessageBox.critical(self, '错误', '没有安装 Adobe PDF Reader')
        self.axWidget.setControl("{233C1507-6A77-46A4-9443-F871F945D258}")
        self.axWidget.dynamicCall(
            'SetVisible (bool Visible)', 'false')  # 不显示窗体
        self.axWidget.dynamicCall('LoadFile(const QString&)',0,  path)
 
    
    
    
    def closeEvent(self, event):
        self.axWidget.close()
        self.axWidget.clear()
        self.layout().removeWidget(self.axWidget)
        del self.axWidget
        super(AxWidget, self).closeEvent(event)